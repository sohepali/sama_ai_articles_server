import hashlib
import json
import mimetypes
import os
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import parse_qsl, urlencode, urljoin, urlparse, urlunparse

import feedparser
import requests
from bs4 import BeautifulSoup
from fastapi import Body, Depends, FastAPI, Header, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel, Field
from sqlalchemy import (
    Boolean,
    Column,
    DateTime,
    ForeignKey,
    Integer,
    LargeBinary,
    String,
    Text,
    UniqueConstraint,
    create_engine,
)
from sqlalchemy.orm import Session, declarative_base, relationship, sessionmaker

try:
    import trafilatura
except Exception:
    trafilatura = None

try:
    import fitz
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None


APP_NAME = "SAMA AI Articles Server"
APP_VERSION = "1.0.0"

DATABASE_URL = os.getenv("DATABASE_URL") or os.getenv("NEWS_DATABASE_URL")
if DATABASE_URL and DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = "postgresql://" + DATABASE_URL[len("postgres://") :]

ADMIN_TOKEN = os.getenv("ADMIN_TOKEN") or os.getenv("NEWS_ADMIN_TOKEN")
APP_SYNC_TOKEN = os.getenv("APP_SYNC_TOKEN") or os.getenv("NEWS_APP_SYNC_TOKEN")

MIN_FULL_CHARS = int(os.getenv("NEWS_MIN_FULL_CHARS", "600"))
MAX_ARTICLES_PER_SOURCE = int(os.getenv("NEWS_MAX_ARTICLES_PER_SOURCE", "25"))
MAX_DOCUMENTS_PER_ARTICLE = int(os.getenv("NEWS_MAX_DOCUMENTS_PER_ARTICLE", "3"))
MAX_DOCUMENT_BYTES = int(os.getenv("NEWS_MAX_DOCUMENT_BYTES", str(20 * 1024 * 1024)))
REQUEST_TIMEOUT = int(os.getenv("NEWS_REQUEST_TIMEOUT", "25"))
APPROVE_PARTIAL = os.getenv("NEWS_APPROVE_PARTIAL", "").strip().lower() in {"1", "true", "yes", "y"}

if not DATABASE_URL:
    raise RuntimeError("DATABASE_URL is required. Use your PostgreSQL connection string.")

engine = create_engine(
    DATABASE_URL,
    pool_pre_ping=True,
    pool_size=5,
    max_overflow=10,
)
SessionLocal = sessionmaker(bind=engine, autocommit=False, autoflush=False)
Base = declarative_base()
bearer_scheme = HTTPBearer(auto_error=False)


class ArticleSource(Base):
    __tablename__ = "sama_article_sources"

    id = Column(Integer, primary_key=True)
    name = Column(String(255), unique=True, index=True, nullable=False)
    source_type = Column(String(50), default="rss")
    rss_url = Column(Text, nullable=True)
    api_url = Column(Text, nullable=True)
    sitemap_url = Column(Text, nullable=True)
    source_page_url = Column(Text, nullable=True)
    base_url = Column(Text, nullable=True)
    country_scope = Column(Text, nullable=True)
    extractor_type = Column(String(80), default="generic")
    active = Column(Boolean, default=True)
    approve_by_default = Column(Boolean, default=True)
    update_interval_minutes = Column(Integer, default=60)
    extra_config = Column(Text, nullable=True)
    last_run_at = Column(DateTime, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    articles = relationship("Article", back_populates="source")


class Article(Base):
    __tablename__ = "sama_articles"
    __table_args__ = (UniqueConstraint("canonical_url", name="uq_sama_articles_canonical_url"),)

    id = Column(Integer, primary_key=True)
    source_id = Column(Integer, ForeignKey("sama_article_sources.id"), nullable=True)
    source_name = Column(String(255), index=True, nullable=False)
    title = Column(Text, nullable=False)
    url = Column(Text, nullable=True)
    canonical_url = Column(Text, nullable=False)
    published_at = Column(DateTime, nullable=True)
    published_text = Column(String(255), nullable=True)
    collected_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), index=True)
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))
    raw_html = Column(Text, nullable=True)
    content = Column(Text, nullable=True)
    summary = Column(Text, nullable=True)
    content_length = Column(Integer, default=0)
    extraction_status = Column(String(50), default="partial", index=True)
    extractor_used = Column(String(80), nullable=True)
    fetch_error = Column(Text, nullable=True)
    country_tags = Column(Text, nullable=True)
    language = Column(String(20), nullable=True)
    approved = Column(Boolean, default=True, index=True)
    metadata_json = Column(Text, nullable=True)

    source = relationship("ArticleSource", back_populates="articles")
    documents = relationship("ArticleDocument", back_populates="article", cascade="all, delete-orphan")


class ArticleDocument(Base):
    __tablename__ = "sama_article_documents"
    __table_args__ = (UniqueConstraint("article_id", "url", name="uq_sama_article_documents_article_url"),)

    id = Column(Integer, primary_key=True)
    article_id = Column(Integer, ForeignKey("sama_articles.id"), nullable=False)
    url = Column(Text, nullable=False)
    filename = Column(Text, nullable=True)
    content_type = Column(String(120), nullable=True)
    file_size = Column(Integer, default=0)
    file_bytes = Column(LargeBinary, nullable=True)
    extracted_text = Column(Text, nullable=True)
    extraction_status = Column(String(50), default="partial")
    error = Column(Text, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    article = relationship("Article", back_populates="documents")


class IngestionRun(Base):
    __tablename__ = "sama_ingestion_runs"

    id = Column(Integer, primary_key=True)
    started_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), index=True)
    finished_at = Column(DateTime, nullable=True)
    source_name = Column(String(255), nullable=True)
    status = Column(String(50), default="running")
    discovered_count = Column(Integer, default=0)
    saved_count = Column(Integer, default=0)
    partial_count = Column(Integer, default=0)
    failed_count = Column(Integer, default=0)
    message = Column(Text, nullable=True)


Base.metadata.create_all(bind=engine)

app = FastAPI(title=APP_NAME, version=APP_VERSION)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in os.getenv("CORS_ORIGINS", "*").split(",")],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class SourceIn(BaseModel):
    name: str = Field(..., min_length=2)
    source_type: str = "rss"
    rss_url: Optional[str] = None
    api_url: Optional[str] = None
    sitemap_url: Optional[str] = None
    source_page_url: Optional[str] = None
    base_url: Optional[str] = None
    country_scope: Optional[str] = None
    extractor_type: str = "generic"
    active: bool = True
    approve_by_default: bool = True
    update_interval_minutes: int = 60
    extra_config: Optional[Dict[str, Any]] = None


class RunIn(BaseModel):
    source_name: Optional[str] = None
    max_articles_per_source: Optional[int] = None


def get_db() -> Session:
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def require_admin(x_admin_token: Optional[str] = Header(None)) -> None:
    if not ADMIN_TOKEN:
        raise HTTPException(status_code=500, detail="ADMIN_TOKEN is not configured on the server.")
    if x_admin_token != ADMIN_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid admin token.")


def require_sync_token(
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(bearer_scheme),
    x_app_token: Optional[str] = Header(None),
) -> None:
    if not APP_SYNC_TOKEN:
        return
    supplied = x_app_token or (credentials.credentials if credentials else None)
    if supplied != APP_SYNC_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid app sync token.")


def utc_now() -> datetime:
    return datetime.now(timezone.utc)


def parse_datetime(value: Any) -> Tuple[Optional[datetime], Optional[str]]:
    if value in (None, ""):
        return None, None
    text = str(value).strip()
    if not text:
        return None, None
    normalized = text.replace("Z", "+00:00")
    try:
        return datetime.fromisoformat(normalized), text
    except Exception:
        pass
    for fmt in ("%Y-%m-%d", "%d %b %Y", "%B %d, %Y", "%Y/%m/%d", "%d/%m/%Y"):
        try:
            return datetime.strptime(text, fmt).replace(tzinfo=timezone.utc), text
        except Exception:
            continue
    return None, text


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        value = "\n".join(str(item) for item in value)
    return re.sub(r"\s+", " ", str(value)).strip()


def canonicalize_url(url: str) -> str:
    parsed = urlparse(url.strip())
    if not parsed.scheme or not parsed.netloc:
        return url.strip()
    query = [
        (key, value)
        for key, value in parse_qsl(parsed.query, keep_blank_values=True)
        if not key.lower().startswith("utm_") and key.lower() not in {"fbclid", "gclid"}
    ]
    return urlunparse(
        (
            parsed.scheme.lower(),
            parsed.netloc.lower(),
            parsed.path.rstrip("/") or "/",
            "",
            urlencode(query, doseq=True),
            "",
        )
    )


def generated_import_url(item: Dict[str, Any], source_name: str) -> str:
    title = clean_text(first_value(item, "title", "Title"))
    date_text = clean_text(first_value(item, "published_at", "published", "Date"))
    digest = hashlib.sha256(f"{source_name}|{title}|{date_text}".encode("utf-8")).hexdigest()[:24]
    return f"import://{source_name.lower().replace(' ', '-')}/{digest}"


def first_value(data: Dict[str, Any], *keys: str) -> Any:
    lowered = {str(key).lower(): value for key, value in data.items()}
    for key in keys:
        if key in data and data[key] not in (None, ""):
            return data[key]
        value = lowered.get(key.lower())
        if value not in (None, ""):
            return value
    return None


def classify_status(content: str, requested_status: Optional[str] = None) -> str:
    requested = (requested_status or "").strip().lower()
    if requested in {"success", "partial", "failed"}:
        return requested
    if not content:
        return "failed"
    if len(content) >= MIN_FULL_CHARS:
        return "success"
    return "partial"


def infer_country_tags(text: str) -> str:
    countries = [
        "Afghanistan",
        "Egypt",
        "Ethiopia",
        "Iran",
        "Iraq",
        "Jordan",
        "Lebanon",
        "Libya",
        "Mali",
        "Myanmar",
        "Palestine",
        "Somalia",
        "South Sudan",
        "Sudan",
        "Syria",
        "Türkiye",
        "Turkey",
        "Ukraine",
        "Yemen",
    ]
    found = []
    lowered = text.lower()
    for country in countries:
        if country.lower() in lowered and country not in found:
            found.append(country)
    return ", ".join(found)


def get_or_create_source(db: Session, source_name: str, source_type: str = "json_import") -> ArticleSource:
    source = db.query(ArticleSource).filter(ArticleSource.name == source_name).first()
    if source:
        return source
    source = ArticleSource(name=source_name, source_type=source_type, extractor_type=source_type)
    db.add(source)
    db.flush()
    return source


def upsert_article_record(
    db: Session,
    source: ArticleSource,
    title: str,
    url: str,
    content: str,
    published_at: Optional[datetime],
    published_text: Optional[str],
    raw_html: Optional[str],
    summary: Optional[str],
    status: str,
    extractor_used: str,
    fetch_error: Optional[str],
    country_tags: Optional[str],
    metadata: Optional[Dict[str, Any]],
) -> Tuple[Article, bool]:
    canonical_url = canonicalize_url(url)
    article = db.query(Article).filter(Article.canonical_url == canonical_url).first()
    created = False
    approved = status == "success" or (status == "partial" and (APPROVE_PARTIAL or source.approve_by_default))
    if not article:
        article = Article(
            source=source,
            source_name=source.name,
            title=title,
            url=url,
            canonical_url=canonical_url,
            collected_at=utc_now(),
        )
        db.add(article)
        created = True

    article.source = source
    article.source_name = source.name
    article.title = title or article.title
    article.url = url or article.url
    article.published_at = published_at or article.published_at
    article.published_text = published_text or article.published_text
    article.raw_html = raw_html or article.raw_html
    article.content = content or article.content
    article.summary = summary or article.summary
    article.content_length = len(article.content or "")
    article.extraction_status = status
    article.extractor_used = extractor_used
    article.fetch_error = fetch_error
    article.country_tags = country_tags or infer_country_tags(f"{title} {content}")
    article.approved = approved
    article.updated_at = utc_now()
    article.metadata_json = json.dumps(metadata or {}, ensure_ascii=False)
    db.flush()
    return article, created


def fetch_html(url: str) -> str:
    response = requests.get(
        url,
        timeout=REQUEST_TIMEOUT,
        headers={
            "User-Agent": "SAMA AI Articles Server/1.0 (+https://samaail.netlify.app)",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        },
    )
    response.raise_for_status()
    response.encoding = response.encoding or "utf-8"
    return response.text


def extract_text_from_html(html: str, url: str = "") -> Tuple[str, str]:
    if trafilatura:
        extracted = trafilatura.extract(html, url=url, include_comments=False, include_tables=True)
        if extracted and len(extracted.strip()) >= 200:
            return clean_text(extracted), "trafilatura"

    soup = BeautifulSoup(html, "html.parser")
    for selector in ["script", "style", "nav", "footer", "header", "aside", "form"]:
        for tag in soup.select(selector):
            tag.decompose()
    article = soup.find("article") or soup.find("main") or soup.body or soup
    paragraphs = [clean_text(p.get_text(" ")) for p in article.find_all(["p", "li", "h2", "h3"])]
    text = " ".join(p for p in paragraphs if p)
    return clean_text(text), "beautifulsoup"


def find_document_links(html: str, base_url: str) -> List[str]:
    soup = BeautifulSoup(html, "html.parser")
    urls = []
    for anchor in soup.find_all("a", href=True):
        href = urljoin(base_url, anchor["href"])
        path = urlparse(href).path.lower()
        if path.endswith((".pdf", ".doc", ".docx")) and href not in urls:
            urls.append(href)
    return urls[:MAX_DOCUMENTS_PER_ARTICLE]


def download_document(url: str) -> Tuple[bytes, str]:
    response = requests.get(
        url,
        timeout=REQUEST_TIMEOUT,
        headers={"User-Agent": "SAMA AI Articles Server/1.0"},
        stream=True,
    )
    response.raise_for_status()
    content_type = response.headers.get("content-type") or mimetypes.guess_type(url)[0] or "application/octet-stream"
    chunks = []
    total = 0
    for chunk in response.iter_content(chunk_size=65536):
        if not chunk:
            continue
        total += len(chunk)
        if total > MAX_DOCUMENT_BYTES:
            raise ValueError("Document is larger than NEWS_MAX_DOCUMENT_BYTES.")
        chunks.append(chunk)
    return b"".join(chunks), content_type


def extract_document_text(file_bytes: bytes, filename: str, content_type: str) -> str:
    lower_name = filename.lower()
    if (lower_name.endswith(".pdf") or "pdf" in content_type) and fitz:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return clean_text("\n".join(page.get_text("text") for page in doc))
    if lower_name.endswith(".docx") and DocxDocument:
        import io

        doc = DocxDocument(io.BytesIO(file_bytes))
        return clean_text("\n".join(paragraph.text for paragraph in doc.paragraphs))
    return ""


def discover_urls(source: ArticleSource) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    if source.rss_url:
        feed = feedparser.parse(source.rss_url)
        for entry in feed.entries[:MAX_ARTICLES_PER_SOURCE]:
            items.append(
                {
                    "title": entry.get("title") or "",
                    "url": entry.get("link") or entry.get("id") or "",
                    "published": entry.get("published") or entry.get("updated") or "",
                    "summary": entry.get("summary") or "",
                }
            )
    elif source.sitemap_url:
        html = fetch_html(source.sitemap_url)
        soup = BeautifulSoup(html, "xml")
        for loc in soup.find_all("loc")[:MAX_ARTICLES_PER_SOURCE]:
            items.append({"title": "", "url": loc.get_text(strip=True), "published": "", "summary": ""})
    elif source.source_page_url:
        html = fetch_html(source.source_page_url)
        soup = BeautifulSoup(html, "html.parser")
        for anchor in soup.find_all("a", href=True):
            href = urljoin(source.source_page_url, anchor["href"])
            title = clean_text(anchor.get_text(" "))
            if title and href.startswith("http"):
                items.append({"title": title, "url": href, "published": "", "summary": ""})
            if len(items) >= MAX_ARTICLES_PER_SOURCE:
                break
    return items


def save_documents_for_article(db: Session, article: Article, html: str, base_url: str) -> None:
    for doc_url in find_document_links(html, base_url):
        existing = (
            db.query(ArticleDocument)
            .filter(ArticleDocument.article_id == article.id, ArticleDocument.url == doc_url)
            .first()
        )
        if existing:
            continue
        document = ArticleDocument(article=article, url=doc_url, filename=os.path.basename(urlparse(doc_url).path))
        try:
            file_bytes, content_type = download_document(doc_url)
            document.content_type = content_type
            document.file_size = len(file_bytes)
            document.file_bytes = file_bytes
            document.extracted_text = extract_document_text(file_bytes, document.filename or "", content_type)
            document.extraction_status = "success" if document.extracted_text else "partial"
        except Exception as exc:
            document.extraction_status = "failed"
            document.error = str(exc)
        db.add(document)


def ingest_source(db: Session, source: ArticleSource, max_articles: int) -> Dict[str, int]:
    stats = {"discovered": 0, "saved": 0, "partial": 0, "failed": 0}
    entries = discover_urls(source)[:max_articles]
    stats["discovered"] = len(entries)

    for entry in entries:
        url = clean_text(entry.get("url"))
        if not url:
            stats["failed"] += 1
            continue
        title = clean_text(entry.get("title")) or url
        published_at, published_text = parse_datetime(entry.get("published"))
        summary = clean_text(entry.get("summary"))
        content = summary
        raw_html = None
        status = "partial"
        extractor_used = "rss"
        error = None
        try:
            raw_html = fetch_html(url)
            extracted, extractor_used = extract_text_from_html(raw_html, url)
            if extracted:
                content = extracted
            status = classify_status(content)
        except Exception as exc:
            error = str(exc)
            status = classify_status(content)

        article, _ = upsert_article_record(
            db=db,
            source=source,
            title=title,
            url=url,
            content=content,
            published_at=published_at,
            published_text=published_text,
            raw_html=raw_html,
            summary=summary,
            status=status,
            extractor_used=extractor_used,
            fetch_error=error,
            country_tags=None,
            metadata=entry,
        )
        if raw_html:
            save_documents_for_article(db, article, raw_html, url)
        if status == "success":
            stats["saved"] += 1
        elif status == "partial":
            stats["partial"] += 1
        else:
            stats["failed"] += 1
    source.last_run_at = utc_now()
    db.commit()
    return stats


def normalize_import_item(item: Dict[str, Any], fallback_source: Optional[str] = None) -> Dict[str, Any]:
    source_name = clean_text(first_value(item, "source", "Source", "source_name")) or fallback_source or "Imported"
    title = clean_text(first_value(item, "title", "Title", "headline")) or "Untitled article"
    url = clean_text(first_value(item, "url", "URL", "link", "original_url", "Original_URL"))
    if not url:
        url = generated_import_url(item, source_name)
    content = clean_text(first_value(item, "content", "Content", "full_text", "article_text", "text"))
    summary = clean_text(first_value(item, "summary", "rss_summary", "description"))
    if not content:
        content = summary
    published_at, published_text = parse_datetime(first_value(item, "published_at", "published", "Date", "date"))
    status = classify_status(content, clean_text(first_value(item, "extraction_status", "Extraction_Status")))
    country_tags = clean_text(first_value(item, "country_tags", "Country_Tags", "countries"))
    return {
        "source_name": source_name,
        "title": title,
        "url": url,
        "content": content,
        "summary": summary,
        "published_at": published_at,
        "published_text": published_text,
        "status": status,
        "country_tags": country_tags,
        "metadata": item,
    }


@app.get("/health")
def health(db: Session = Depends(get_db)) -> Dict[str, Any]:
    article_count = db.query(Article).count()
    source_count = db.query(ArticleSource).count()
    return {
        "ok": True,
        "app": APP_NAME,
        "version": APP_VERSION,
        "sources": source_count,
        "articles": article_count,
        "time": utc_now().isoformat(),
    }


@app.get("/admin/dashboard", response_class=HTMLResponse)
def dashboard(_: None = Depends(require_admin), db: Session = Depends(get_db)) -> str:
    sources = db.query(ArticleSource).order_by(ArticleSource.name).all()
    articles = db.query(Article).order_by(Article.updated_at.desc()).limit(20).all()
    source_rows = "".join(
        f"<tr><td>{s.name}</td><td>{s.source_type}</td><td>{'Yes' if s.active else 'No'}</td><td>{s.last_run_at or ''}</td></tr>"
        for s in sources
    )
    article_rows = "".join(
        f"<tr><td>{a.source_name}</td><td>{a.title}</td><td>{a.extraction_status}</td><td>{a.content_length}</td></tr>"
        for a in articles
    )
    return f"""
    <!doctype html>
    <html>
    <head>
      <title>SAMA AI Articles Server</title>
      <style>
        body {{ font-family: Arial, sans-serif; margin: 32px; color: #182735; }}
        h1 {{ color: #009edb; }}
        h2 {{ margin-top: 28px; color: #f58220; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 12px; }}
        th, td {{ border-bottom: 1px solid #d8e0e6; padding: 9px; text-align: left; vertical-align: top; }}
        th {{ background: #edf6fb; }}
        .pill {{ display: inline-block; background: #f58220; color: white; padding: 4px 8px; border-radius: 4px; }}
      </style>
    </head>
    <body>
      <h1>SAMA AI Articles Server <span class="pill">{APP_VERSION}</span></h1>
      <p>Standalone article database and ingestion API.</p>
      <h2>Sources</h2>
      <table><tr><th>Name</th><th>Type</th><th>Active</th><th>Last run</th></tr>{source_rows}</table>
      <h2>Latest Articles</h2>
      <table><tr><th>Source</th><th>Title</th><th>Status</th><th>Characters</th></tr>{article_rows}</table>
    </body>
    </html>
    """


@app.post("/admin/sources")
def upsert_source(payload: SourceIn, _: None = Depends(require_admin), db: Session = Depends(get_db)) -> Dict[str, Any]:
    source = db.query(ArticleSource).filter(ArticleSource.name == payload.name).first()
    if not source:
        source = ArticleSource(name=payload.name)
        db.add(source)
    for field in (
        "source_type",
        "rss_url",
        "api_url",
        "sitemap_url",
        "source_page_url",
        "base_url",
        "country_scope",
        "extractor_type",
        "active",
        "approve_by_default",
        "update_interval_minutes",
    ):
        setattr(source, field, getattr(payload, field))
    source.extra_config = json.dumps(payload.extra_config or {}, ensure_ascii=False)
    source.updated_at = utc_now()
    db.commit()
    return {"ok": True, "source_id": source.id, "name": source.name}


@app.get("/admin/sources")
def list_sources(_: None = Depends(require_admin), db: Session = Depends(get_db)) -> Dict[str, Any]:
    sources = db.query(ArticleSource).order_by(ArticleSource.name).all()
    return {
        "sources": [
            {
                "id": source.id,
                "name": source.name,
                "source_type": source.source_type,
                "rss_url": source.rss_url,
                "sitemap_url": source.sitemap_url,
                "source_page_url": source.source_page_url,
                "active": source.active,
                "last_run_at": source.last_run_at.isoformat() if source.last_run_at else None,
            }
            for source in sources
        ]
    }


@app.post("/admin/sources/seed")
def seed_sources(_: None = Depends(require_admin), db: Session = Depends(get_db)) -> Dict[str, Any]:
    defaults = [
        SourceIn(name="UN News", source_type="rss", rss_url="https://news.un.org/feed/subscribe/en/news/all/rss.xml"),
        SourceIn(name="ReliefWeb", source_type="rss", rss_url="https://reliefweb.int/updates/rss.xml"),
        SourceIn(name="BBC World", source_type="rss", rss_url="https://feeds.bbci.co.uk/news/world/rss.xml"),
        SourceIn(name="Al Jazeera", source_type="rss", rss_url="https://www.aljazeera.com/xml/rss/all.xml"),
        SourceIn(name="The New Humanitarian", source_type="rss", rss_url="https://www.thenewhumanitarian.org/rss.xml"),
    ]
    created = []
    for source_in in defaults:
        existing = db.query(ArticleSource).filter(ArticleSource.name == source_in.name).first()
        if existing:
            continue
        source = ArticleSource(
            name=source_in.name,
            source_type=source_in.source_type,
            rss_url=source_in.rss_url,
            extractor_type="generic",
            active=True,
            approve_by_default=True,
        )
        db.add(source)
        created.append(source_in.name)
    db.commit()
    return {"ok": True, "created": created}


@app.post("/admin/ingest/run")
def run_ingestion(payload: RunIn, _: None = Depends(require_admin), db: Session = Depends(get_db)) -> Dict[str, Any]:
    query = db.query(ArticleSource).filter(ArticleSource.active.is_(True))
    if payload.source_name:
        query = query.filter(ArticleSource.name == payload.source_name)
    sources = query.order_by(ArticleSource.name).all()
    if not sources:
        raise HTTPException(status_code=404, detail="No active sources found.")

    results = []
    for source in sources:
        run = IngestionRun(source_name=source.name)
        db.add(run)
        db.commit()
        try:
            stats = ingest_source(db, source, payload.max_articles_per_source or MAX_ARTICLES_PER_SOURCE)
            run.status = "success"
            run.discovered_count = stats["discovered"]
            run.saved_count = stats["saved"]
            run.partial_count = stats["partial"]
            run.failed_count = stats["failed"]
            run.message = json.dumps(stats)
            results.append({"source": source.name, **stats})
        except Exception as exc:
            db.rollback()
            run = db.query(IngestionRun).filter(IngestionRun.id == run.id).first() or IngestionRun(source_name=source.name)
            db.add(run)
            run.status = "failed"
            run.message = str(exc)
            results.append({"source": source.name, "error": str(exc)})
        finally:
            run.finished_at = utc_now()
            db.commit()
    return {"ok": True, "results": results}


@app.post("/admin/articles/import-json")
@app.post("/admin/news/import-json")
def import_json_articles(
    payload: Any = Body(...),
    _: None = Depends(require_admin),
    db: Session = Depends(get_db),
) -> Dict[str, Any]:
    if isinstance(payload, list):
        raw_items = payload
        fallback_source = None
    elif isinstance(payload, dict):
        raw_items = payload.get("articles") or payload.get("items") or [payload]
        fallback_source = clean_text(payload.get("source_name") or payload.get("source"))
    else:
        raise HTTPException(status_code=400, detail="Expected a JSON list or an object with articles/items.")

    if not isinstance(raw_items, list):
        raise HTTPException(status_code=400, detail="articles/items must be a list.")

    saved = 0
    updated = 0
    failed = []
    for index, raw_item in enumerate(raw_items, start=1):
        if not isinstance(raw_item, dict):
            failed.append({"index": index, "reason": "Item is not a JSON object."})
            continue
        try:
            item = normalize_import_item(raw_item, fallback_source)
            source = get_or_create_source(db, item["source_name"], "json_import")
            article, created = upsert_article_record(
                db=db,
                source=source,
                title=item["title"],
                url=item["url"],
                content=item["content"],
                published_at=item["published_at"],
                published_text=item["published_text"],
                raw_html=clean_text(first_value(raw_item, "raw_html", "html")) or None,
                summary=item["summary"],
                status=item["status"],
                extractor_used="json_import",
                fetch_error=None if not item["url"].startswith("import://") else "Original URL was not provided in the import payload.",
                country_tags=item["country_tags"],
                metadata=item["metadata"],
            )
            saved += 1 if created else 0
            updated += 0 if created else 1
        except Exception as exc:
            failed.append({"index": index, "reason": str(exc)})
    db.commit()
    return {"ok": True, "received": len(raw_items), "created": saved, "updated": updated, "failed": failed}


@app.get("/admin/articles")
def list_admin_articles(
    _: None = Depends(require_admin),
    db: Session = Depends(get_db),
    limit: int = Query(100, ge=1, le=500),
    status: Optional[str] = Query(None),
    source: Optional[str] = Query(None),
) -> Dict[str, Any]:
    query = db.query(Article)
    if status:
        query = query.filter(Article.extraction_status == status)
    if source:
        query = query.filter(Article.source_name == source)
    articles = query.order_by(Article.updated_at.desc()).limit(limit).all()
    return {"articles": [article_to_api(article, include_content=False) for article in articles]}


@app.get("/admin/runs")
def list_runs(_: None = Depends(require_admin), db: Session = Depends(get_db), limit: int = Query(50, ge=1, le=200)) -> Dict[str, Any]:
    runs = db.query(IngestionRun).order_by(IngestionRun.started_at.desc()).limit(limit).all()
    return {
        "runs": [
            {
                "id": run.id,
                "source_name": run.source_name,
                "status": run.status,
                "started_at": run.started_at.isoformat() if run.started_at else None,
                "finished_at": run.finished_at.isoformat() if run.finished_at else None,
                "discovered_count": run.discovered_count,
                "saved_count": run.saved_count,
                "partial_count": run.partial_count,
                "failed_count": run.failed_count,
                "message": run.message,
            }
            for run in runs
        ]
    }


def article_to_api(article: Article, include_content: bool = True) -> Dict[str, Any]:
    data = {
        "id": article.id,
        "Source": article.source_name,
        "Title": article.title,
        "Date": article.published_at.date().isoformat() if article.published_at else article.published_text,
        "URL": article.url,
        "Extraction_Status": article.extraction_status,
        "Content_Length": article.content_length,
        "Country_Tags": article.country_tags,
        "Updated_At": article.updated_at.isoformat() if article.updated_at else None,
        "Scraped_At": article.collected_at.isoformat() if article.collected_at else None,
        "Document_Count": len(article.documents or []),
    }
    if include_content:
        data["Content"] = article.content or ""
        data["Summary"] = article.summary or ""
    return data


@app.get("/articles/sync")
def sync_articles(
    _: None = Depends(require_sync_token),
    db: Session = Depends(get_db),
    since: Optional[str] = Query(None),
    limit: int = Query(500, ge=1, le=2000),
    source: Optional[str] = Query(None),
    country: Optional[str] = Query(None),
) -> Dict[str, Any]:
    query = db.query(Article).filter(Article.approved.is_(True))
    if since:
        since_dt, _ = parse_datetime(since)
        if since_dt:
            query = query.filter(Article.updated_at > since_dt)
    if source:
        query = query.filter(Article.source_name == source)
    if country:
        query = query.filter(Article.country_tags.ilike(f"%{country}%"))
    articles = query.order_by(Article.updated_at.desc()).limit(limit).all()
    return {
        "ok": True,
        "count": len(articles),
        "server_time": utc_now().isoformat(),
        "articles": [article_to_api(article) for article in articles],
    }


@app.get("/articles/{article_id}")
def get_article(article_id: int, _: None = Depends(require_sync_token), db: Session = Depends(get_db)) -> Dict[str, Any]:
    article = db.query(Article).filter(Article.id == article_id, Article.approved.is_(True)).first()
    if not article:
        raise HTTPException(status_code=404, detail="Article not found.")
    data = article_to_api(article)
    data["documents"] = [
        {
            "id": document.id,
            "url": document.url,
            "filename": document.filename,
            "content_type": document.content_type,
            "file_size": document.file_size,
            "extraction_status": document.extraction_status,
            "extracted_text": document.extracted_text,
        }
        for document in article.documents
    ]
    return data
